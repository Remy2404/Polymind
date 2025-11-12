import logging
import io
import re
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import ContextTypes

try:
    from spire.doc import Document, Section, FileFormat
    from spire.doc.common import (
        Color,
        HorizontalAlignment,
        VerticalAlignment,
        BorderStyle,
        HyperlinkType,
    )

    SPIRE_AVAILABLE = True
except ImportError:
    SPIRE_AVAILABLE = False
    Document = None
    Section = None
    FileFormat = None
    Color = None
    HorizontalAlignment = None
    VerticalAlignment = None
    BorderStyle = None
    HyperlinkType = None

# Try to import python-docx as fallback
try:
    from docx import Document as DocxDocument

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    DocxDocument = None
from pydantic import BaseModel, Field


class ExportContent(BaseModel):
    """Simplified export content structure - content only"""

    content: str = Field(..., description="Main document content")


class SpireDocumentExporter:
    """Simplified document exporter using Spire.Doc"""

    def __init__(self, logger: logging.Logger):
        self.logger = logger
        self.current_table = None
        self.table_rows = []
        self.in_table = False

    def create_docx(self, content: str) -> bytes:
        """Create DOCX using Spire.Doc - simplified to export content only"""
        if not SPIRE_AVAILABLE:
            raise ImportError(
                "Spire.Doc is not available. Please install it to use DOCX export functionality."
            )
        try:
            document = Document()
            try:
                document.BuiltinDocumentProperties.Title = ""
                document.BuiltinDocumentProperties.Subject = ""
                document.BuiltinDocumentProperties.Author = ""
                document.BuiltinDocumentProperties.Company = ""
                document.BuiltinDocumentProperties.Comments = ""
                document.BuiltinDocumentProperties.Keywords = ""
                document.BuiltinDocumentProperties.Category = ""
                document.BuiltinDocumentProperties.Manager = ""
            except Exception:
                pass
            section = document.AddSection()
            self._add_formatted_content(section, content)
            import tempfile
            import os

            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
                temp_path = temp_file.name
            try:
                document.SaveToFile(temp_path, FileFormat.Docx)
                document.Close()
                with open(temp_path, "rb") as f:
                    document_bytes = f.read()
                return document_bytes
            finally:
                os.unlink(temp_path)
        except Exception as e:
            self.logger.error(f"Spire.Doc export failed: {e}")
            raise

    def _add_formatted_content(self, section, content: str):
        """Enhanced markdown to DOCX formatting with comprehensive support"""
        lines = content.split("\n")
        self.current_table = None
        self.table_rows = []
        self.in_table = False
        for line in lines:
            line = line.strip()
            if not line:
                if self.in_table:
                    self._finalize_current_table(section)
                section.AddParagraph()
                continue
            if any(
                phrase in line.lower()
                for phrase in [
                    "evaluation warning",
                    "spire.doc for python",
                    "document was created with spire.doc",
                ]
            ):
                continue
            paragraph = section.AddParagraph()
            if line.startswith("# "):
                text = line[2:].strip()
                text_range = paragraph.AppendText(text)
                text_range.CharacterFormat.Bold = True
                text_range.CharacterFormat.FontSize = 18
                text_range.CharacterFormat.FontName = "Arial"
                text_range.CharacterFormat.TextColor = Color.FromArgb(0, 32, 96, 148)
                paragraph.Format.SpaceAfter = 15
                paragraph.Format.SpaceBefore = 10
            elif line.startswith("## "):
                text = line[3:].strip()
                text_range = paragraph.AppendText(text)
                text_range.CharacterFormat.Bold = True
                text_range.CharacterFormat.FontSize = 16
                text_range.CharacterFormat.FontName = "Arial"
                text_range.CharacterFormat.TextColor = Color.FromArgb(0, 68, 114, 196)
                paragraph.Format.SpaceAfter = 12
                paragraph.Format.SpaceBefore = 8
            elif line.startswith("### "):
                text = line[4:].strip()
                text_range = paragraph.AppendText(text)
                text_range.CharacterFormat.Bold = True
                text_range.CharacterFormat.FontSize = 14
                text_range.CharacterFormat.FontName = "Arial"
                text_range.CharacterFormat.TextColor = Color.FromArgb(0, 84, 130, 53)
                paragraph.Format.SpaceAfter = 10
                paragraph.Format.SpaceBefore = 6
            elif (
                line.startswith("• ") or line.startswith("- ") or line.startswith("* ")
            ):
                bullet_char = "•"
                text = line[2:].strip()
                self._add_formatted_text_with_inline(paragraph, f"{bullet_char} {text}")
                paragraph.Format.LeftIndent = 20
                paragraph.Format.FirstLineIndent = -10
                paragraph.Format.SpaceAfter = 3
            elif re.match(r"^\d+\.\s", line):
                paragraph.Format.LeftIndent = 20
                paragraph.Format.FirstLineIndent = -10
                paragraph.Format.SpaceAfter = 3
            elif line.startswith("> "):
                text = line[2:].strip()
                text_range = paragraph.AppendText(f"❝ {text}")
                text_range.CharacterFormat.FontName = "Arial"
                text_range.CharacterFormat.Italic = True
                text_range.CharacterFormat.TextColor = Color.FromArgb(0, 128, 128, 128)
                paragraph.Format.LeftIndent = 25
                paragraph.Format.RightIndent = 25
                paragraph.Format.SpaceBefore = 6
                paragraph.Format.SpaceAfter = 6
            elif line.startswith("---") or line.startswith("___"):
                text_range = paragraph.AppendText("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
                text_range.CharacterFormat.FontName = "Arial"
                text_range.CharacterFormat.TextColor = Color.FromArgb(0, 200, 200, 200)
                paragraph.Format.Alignment = HorizontalAlignment.Center
                paragraph.Format.SpaceBefore = 12
                paragraph.Format.SpaceAfter = 12
            elif (
                "|" in line
                and line.count("|") >= 2
                and re.match(r"^\s*\|.*\|\s*$", line)
            ):
                self._collect_table_row(line)
                continue
            else:
                if self.in_table:
                    self._finalize_current_table(section)
                paragraph = section.AddParagraph()
                self._add_formatted_text_with_inline(paragraph, line)
                paragraph.Format.SpaceAfter = 3
        if self.in_table:
            self._finalize_current_table(section)

    def _add_formatted_text_with_inline(self, paragraph, text):
        """Add text with inline formatting support (bold, italic, code, strikethrough, links)"""
        if not text:
            return
        patterns = [
            (r"\*\*(.*?)\*\*", "bold"),
            (r"__(.*?)__", "bold"),
            (r"\*(.*?)\*", "italic"),
            (r"_(.*?)_", "italic"),
            (r"`(.*?)`", "code"),
            (r"~~(.*?)~~", "strikethrough"),
            (r"\[([^\]]+)\]\(([^)]+)\)", "link"),
        ]
        remaining_text = text
        while remaining_text:
            earliest_match = None
            earliest_pos = len(remaining_text)
            earliest_format = None
            pass
            for pattern, format_type in patterns:
                match = re.search(pattern, remaining_text)
                if match and match.start() < earliest_pos:
                    earliest_match = match
                    earliest_pos = match.start()
                    earliest_format = format_type
                    pass
            if earliest_match:
                if earliest_pos > 0:
                    before_text = remaining_text[:earliest_pos]
                    if before_text.strip():
                        text_range = paragraph.AppendText(before_text)
                        text_range.CharacterFormat.FontName = "Arial"
                if earliest_format == "link":
                    link_text = earliest_match.group(1)
                    link_url = earliest_match.group(2)
                    hyperlink = paragraph.AppendHyperlink(
                        link_url, link_text, HyperlinkType.WebLink
                    )
                    hyperlink.CharacterFormat.FontName = "Arial"
                    hyperlink.CharacterFormat.TextColor = Color.FromArgb(
                        0, 17, 100, 180
                    )
                else:
                    formatted_text = earliest_match.group(1)
                    text_range = paragraph.AppendText(formatted_text)
                    text_range.CharacterFormat.FontName = "Arial"
                    if earliest_format == "bold":
                        text_range.CharacterFormat.Bold = True
                    elif earliest_format == "italic":
                        text_range.CharacterFormat.Italic = True
                    elif earliest_format == "code":
                        text_range.CharacterFormat.FontName = "Courier New"
                        text_range.CharacterFormat.FontSize = 10
                        text_range.CharacterFormat.TextColor = Color.FromArgb(
                            0, 196, 26, 22
                        )
                    elif earliest_format == "strikethrough":
                        text_range.CharacterFormat.StrikeThrough = True
                        text_range.CharacterFormat.TextColor = Color.FromArgb(
                            0, 128, 128, 128
                        )
                remaining_text = remaining_text[earliest_match.end() :]
            else:
                if remaining_text.strip():
                    text_range = paragraph.AppendText(remaining_text)
                    text_range.CharacterFormat.FontName = "Arial"
                break

    def _collect_table_row(self, line):
        """Collect table row data for proper table creation"""
        if not re.match(r"^\s*\|.*\|\s*$", line):
            return
        raw_cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r"^[-:]+$", cell) for cell in raw_cells):
            return
        cells = [
            (c[0].upper() + c[1:] if c and c[0].islower() else c) for c in raw_cells
        ]
        if not self.in_table:
            self.in_table = True
            self.table_rows = []
        self.table_rows.append(cells)

    def _finalize_current_table(self, section):
        """Create actual DOCX table from collected table rows - improved version"""
        if not self.in_table or not self.table_rows:
            return
        try:
            max_cols = (
                max(len(row) for row in self.table_rows) if self.table_rows else 0
            )
            if max_cols == 0:
                return
            table = section.AddTable(True)
            table.ResetCells(len(self.table_rows), max_cols)
            for row_idx, row_data in enumerate(self.table_rows):
                if row_idx < len(table.Rows):
                    table_row = table.Rows[row_idx]
                    try:
                        table_row.Height = 25 if row_idx == 0 else 20
                    except Exception:
                        pass
                    if row_idx == 0:
                        try:
                            table_row.IsHeader = True
                            try:
                                table_row.RowFormat.BackColor = Color.get_LightBlue()
                            except AttributeError:
                                try:
                                    table_row.RowFormat.BackColor = Color.FromArgb(
                                        240, 248, 255
                                    )
                                except Exception:
                                    pass
                        except Exception:
                            pass
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < len(table_row.Cells):
                            cell = table_row.Cells[col_idx]
                            try:
                                cell_paragraph = cell.AddParagraph()
                                text_range = cell_paragraph.AppendText(str(cell_text))
                                text_range.CharacterFormat.FontName = "Arial"
                                text_range.CharacterFormat.FontSize = (
                                    10 if row_idx > 0 else 11
                                )
                                if row_idx == 0:
                                    text_range.CharacterFormat.Bold = True
                                    try:
                                        cell_paragraph.Format.HorizontalAlignment = (
                                            HorizontalAlignment.Center
                                        )
                                    except Exception:
                                        pass
                                else:
                                    try:
                                        cell_paragraph.Format.HorizontalAlignment = (
                                            HorizontalAlignment.Left
                                        )
                                    except Exception:
                                        pass
                                try:
                                    cell.CellFormat.VerticalAlignment = (
                                        VerticalAlignment.Middle
                                    )
                                except Exception:
                                    pass
                            except Exception as cell_error:
                                self.logger.warning(
                                    f"Error formatting cell [{row_idx}][{col_idx}]: {cell_error}"
                                )
            self._apply_table_borders(table)
        except Exception as e:
            self.logger.error(f"Error creating table: {e}")
            self._create_text_fallback_table(section)
        finally:
            self.in_table = False
            self.table_rows = []
            self.current_table = None

    def _create_text_fallback_table(self, section):
        """Create a well-formatted text table as fallback"""
        if not self.table_rows:
            return
        col_widths = []
        max_cols = max(len(row) for row in self.table_rows)
        for col_idx in range(max_cols):
            max_width = 0
            for row in self.table_rows:
                if col_idx < len(row):
                    max_width = max(max_width, len(str(row[col_idx])))
            col_widths.append(max(max_width, 8))
        for row_idx, row in enumerate(self.table_rows):
            paragraph = section.AddParagraph()
            formatted_cells = []
            for col_idx, cell in enumerate(row):
                if col_idx < len(col_widths):
                    padded_cell = str(cell).ljust(col_widths[col_idx])
                    formatted_cells.append(padded_cell)
            table_text = "| " + " | ".join(formatted_cells) + " |"
            text_range = paragraph.AppendText(table_text)
            text_range.CharacterFormat.FontName = "Courier New"
            text_range.CharacterFormat.FontSize = 10
            if row_idx == 0:
                text_range.CharacterFormat.Bold = True
                separator_paragraph = section.AddParagraph()
                separator_text = (
                    "|" + "|".join(["-" * (w + 2) for w in col_widths]) + "|"
                )
                sep_range = separator_paragraph.AppendText(separator_text)
                sep_range.CharacterFormat.FontName = "Courier New"
                sep_range.CharacterFormat.FontSize = 10
            paragraph.Format.SpaceBefore = 2
            paragraph.Format.SpaceAfter = 2
        section.AddParagraph()

    def _apply_table_borders(self, table):
        """Apply professional borders to the table - fixed version"""
        try:
            for row_idx, row in enumerate(table.Rows):
                for col_idx, cell in enumerate(row.Cells):
                    try:
                        cell_format = cell.CellFormat
                        try:
                            cell_format.Borders.BorderType = BorderStyle.Single
                            cell_format.Borders.LineWidth = 0.5
                            cell_format.Borders.Color = Color.Black
                        except AttributeError:
                            try:
                                borders = cell_format.Borders
                                for border_type in ["Top", "Bottom", "Left", "Right"]:
                                    if hasattr(borders, border_type):
                                        border = getattr(borders, border_type)
                                        if hasattr(border, "BorderType"):
                                            border.BorderType = BorderStyle.Single
                                        if hasattr(border, "LineWidth"):
                                            border.LineWidth = 0.5
                            except Exception:
                                pass
                        except Exception:
                            continue
                    except Exception as cell_error:
                        self.logger.debug(
                            f"Could not format cell [{row_idx}][{col_idx}]: {cell_error}"
                        )
                        continue
        except Exception as e:
            self.logger.warning(f"Could not apply table borders: {e}")

    def _add_emoji_support(self, text_range):
        """Enhance emoji rendering with proper font support"""
        text_range.CharacterFormat.FontName = "Segoe UI Emoji"
        return text_range


class DocxDocumentExporter:
    """Fallback document exporter using python-docx"""

    def __init__(self, logger: logging.Logger):
        self.logger = logger

    def create_docx(self, content: str) -> bytes:
        """Create DOCX using python-docx as fallback"""
        if not DOCX_AVAILABLE:
            raise ImportError(
                "Neither Spire.Doc nor python-docx is available. Please install document export dependencies."
            )

        try:
            document = DocxDocument()

            # Split content into lines and process
            lines = content.split("\n")

            for line in lines:
                line = line.strip()
                if not line:
                    continue

                if line.startswith("# "):
                    # Heading 1
                    document.add_heading(line[2:].strip(), level=1)
                elif line.startswith("## "):
                    # Heading 2
                    document.add_heading(line[3:].strip(), level=2)
                elif line.startswith("### "):
                    # Heading 3
                    document.add_heading(line[4:].strip(), level=3)
                elif line.startswith("- ") or line.startswith("* "):
                    # Bullet point
                    document.add_paragraph(line[2:].strip(), style="List Bullet")
                elif re.match(r"^\d+\.\s", line):
                    # Numbered list
                    document.add_paragraph(line, style="List Number")
                else:
                    # Regular paragraph
                    document.add_paragraph(line)

            # Save to bytes
            import io

            docx_bytes = io.BytesIO()
            document.save(docx_bytes)
            docx_bytes.seek(0)
            return docx_bytes.getvalue()

        except Exception as e:
            self.logger.error(f"python-docx export failed: {e}")
            raise

    def _process_inline_formatting(self, paragraph, text):
        """Process inline markdown formatting like **bold**, *italic*, etc."""
        self._add_formatted_text_with_inline(paragraph, text)


class EnhancedExportCommands:
    """Streamlined export commands using Spire.Doc with python-docx fallback"""

    def __init__(self, gemini_api, user_data_manager, telegram_logger):
        self.gemini_api = gemini_api
        self.user_data_manager = user_data_manager
        self.telegram_logger = telegram_logger
        self.logger = logging.getLogger(__name__)

        # Initialize document exporters with fallback
        if SPIRE_AVAILABLE:
            self.document_exporter = SpireDocumentExporter(self.logger)
            self.logger.info("Using Spire.Doc for document export")
        elif DOCX_AVAILABLE:
            self.document_exporter = DocxDocumentExporter(self.logger)
            self.logger.info("Using python-docx fallback for document export")
        else:
            self.document_exporter = None
            self.logger.warning("No document export library available")

        self.memory_manager = self._init_memory_manager()

    def _init_memory_manager(self):
        """Initialize memory manager with fallback"""
        try:
            from src.services.memory_context.memory_manager import MemoryManager
            from src.database.connection import get_database

            db, _ = get_database()
            return MemoryManager(db=db)
        except Exception as e:
            self.logger.warning(f"Memory manager initialization failed: {e}")
            return None

    async def export_to_document(
        self, update: Update, context: ContextTypes.DEFAULT_TYPE
    ) -> None:
        """Main export entry point"""
        try:
            user_id = update.effective_user.id
            self.telegram_logger.log_message(
                f"Document export requested by user {user_id}", user_id
            )
            if context.args:
                custom_text = " ".join(context.args)
                context.user_data["doc_export_text"] = custom_text
                await self._show_format_selection(update, context)
                return
            keyboard = [
                [
                    InlineKeyboardButton(
                        "💬 Export Conversation", callback_data="export_conversation"
                    ),
                    InlineKeyboardButton(
                        "✏️ Export Custom Text", callback_data="export_custom"
                    ),
                ],
                [InlineKeyboardButton("❌ Cancel", callback_data="export_cancel")],
            ]
            await update.message.reply_text(
                "📄 **Document Export Options**\n\nChoose what you'd like to export:",
                reply_markup=InlineKeyboardMarkup(keyboard),
                parse_mode="Markdown",
            )
        except Exception as e:
            self.logger.error(f"Error in export_to_document: {e}")
            await update.message.reply_text(
                "❌ An error occurred while processing your export request."
            )

    async def handle_export_callback(
        self, update: Update, context: ContextTypes.DEFAULT_TYPE
    ) -> None:
        """Handle export callbacks"""
        callback_data = update.callback_query.data
        try:
            if callback_data == "export_conversation":
                await self._handle_conversation_export(update, context)
            elif callback_data == "export_custom":
                await self._handle_custom_export(update, context)
            elif callback_data == "export_cancel":
                await update.callback_query.edit_message_text("❌ Export cancelled.")
            elif callback_data.startswith("export_format_"):
                format_type = callback_data.replace("export_format_", "")
                await self._generate_document(update, context, format_type)
        except Exception as e:
            self.logger.error(f"Error handling export callback: {e}")
            await update.callback_query.edit_message_text(
                "❌ An error occurred while processing your request."
            )

    async def handle_export_conversation(
        self, update: Update, context: ContextTypes.DEFAULT_TYPE
    ) -> None:
        """Public method for direct conversation export"""
        try:
            user_id = update.effective_user.id
            conversation_content = await self._get_conversation_history(user_id)
            if not conversation_content:
                error_msg = "❌ No conversation history found to export."
                if update.callback_query:
                    await update.callback_query.edit_message_text(error_msg)
                else:
                    await update.message.reply_text(error_msg)
                return
            context.user_data["doc_export_text"] = conversation_content
            context.user_data["export_type"] = "conversation"
            await self._show_format_selection(update, context)
        except Exception as e:
            self.logger.error(f"Error in handle_export_conversation: {e}")
            error_msg = "❌ Error retrieving conversation history for export."
            if update.callback_query:
                await update.callback_query.edit_message_text(error_msg)
            else:
                await update.message.reply_text(error_msg)

    async def _handle_conversation_export(
        self, update: Update, context: ContextTypes.DEFAULT_TYPE
    ) -> None:
        """Handle conversation export callback"""
        user_id = update.effective_user.id
        conversation_content = await self._get_conversation_history(user_id)
        if not conversation_content:
            await update.callback_query.edit_message_text(
                "❌ No conversation history found to export."
            )
            return
        context.user_data["doc_export_text"] = conversation_content
        context.user_data["export_type"] = "conversation"
        await self._show_format_selection(update, context)

    async def _handle_custom_export(
        self, update: Update, context: ContextTypes.DEFAULT_TYPE
    ) -> None:
        """Handle custom text export"""
        context.user_data["awaiting_doc_text"] = True
        context.user_data["export_type"] = "custom"
        await update.callback_query.edit_message_text(
            "✏️ **Send Custom Text for Export**\n\n"
            "Please send me the text you want to convert to a document.\n"
            "After sending your text, I'll ask you to choose the format (DOCX).",
            parse_mode="Markdown",
        )

    async def _show_format_selection(
        self, update: Update, context: ContextTypes.DEFAULT_TYPE
    ) -> None:
        """Show format selection"""
        keyboard = [
            [InlineKeyboardButton("📝 DOCX Format", callback_data="export_format_docx")]
        ]
        message_text = (
            "📋 **Choose Document Format**\n\nSelect DOCX format for your export:"
        )
        if update.callback_query:
            await update.callback_query.edit_message_text(
                message_text,
                reply_markup=InlineKeyboardMarkup(keyboard),
                parse_mode="Markdown",
            )
        else:
            await update.message.reply_text(
                message_text,
                reply_markup=InlineKeyboardMarkup(keyboard),
                parse_mode="Markdown",
            )

    async def generate_document(
        self, update: Update, context: ContextTypes.DEFAULT_TYPE, format_type: str
    ) -> None:
        """Public method to generate document - delegates to private _generate_document"""
        await self._generate_document(update, context, format_type)

    async def _generate_document(
        self, update: Update, context: ContextTypes.DEFAULT_TYPE, format_type: str
    ) -> None:
        """Generate document using available exporter - with fallback support"""
        try:
            if not self.document_exporter:
                await update.callback_query.edit_message_text(
                    "❌ Document export is not available. Please install Spire.Doc or python-docx."
                )
                return

            content = context.user_data.get("doc_export_text", "")
            if not content.strip():
                await update.callback_query.edit_message_text(
                    "❌ No content available for export."
                )
                return

            await update.callback_query.edit_message_text(
                f"🔄 Generating {format_type.upper()} document... This may take a moment."
            )

            if format_type.lower() == "docx":
                document_bytes = self.document_exporter.create_docx(content)
                await self._send_document(update, context, document_bytes)
            else:
                await update.callback_query.edit_message_text(
                    f"❌ {format_type.upper()} export is not yet implemented."
                )

            self._cleanup_user_data(context)
        except Exception as e:
            self.logger.error(f"Error generating document: {e}")
            error_msg = (
                f"❌ Sorry, there was an error generating your document: {str(e)}"
            )
            if "Spire.Doc is not available" in str(e):
                error_msg += "\n\n💡 Using fallback document generator..."
                # Try with fallback if Spire.Doc failed
                if DOCX_AVAILABLE and not SPIRE_AVAILABLE:
                    try:
                        self.document_exporter = DocxDocumentExporter(self.logger)
                        document_bytes = self.document_exporter.create_docx(content)
                        await self._send_document(update, context, document_bytes)
                        self._cleanup_user_data(context)
                        return
                    except Exception as fallback_e:
                        error_msg = f"❌ Fallback also failed: {str(fallback_e)}"

            await update.callback_query.edit_message_text(error_msg)

    async def _send_document(
        self, update: Update, context: ContextTypes.DEFAULT_TYPE, document_bytes: bytes
    ) -> None:
        """Send document to user - simplified without custom filename"""
        try:
            with io.BytesIO(document_bytes) as doc_io:
                doc_io.name = "export.docx"
                caption = "📄 **Document Export Complete**"
                await context.bot.send_document(
                    chat_id=update.effective_chat.id,
                    document=doc_io,
                    filename="export.docx",
                    caption=caption,
                    parse_mode="Markdown",
                )
            self.telegram_logger.log_message(
                "Document generated and sent successfully", update.effective_user.id
            )
        except Exception as e:
            self.logger.error(f"Error sending document: {e}")
            raise

    def _cleanup_user_data(self, context: ContextTypes.DEFAULT_TYPE) -> None:
        """Clean up user data"""
        keys_to_remove = ["doc_export_text", "export_type", "awaiting_doc_text"]
        for key in keys_to_remove:
            context.user_data.pop(key, None)

    async def _get_conversation_history(self, user_id: int) -> str:
        """Get conversation history from various sources, omitting unwanted headers/messages."""
        try:
            formatted_content = ""
            has_content = False
            if self.memory_manager:
                try:
                    conversation_id = f"user_{user_id}"
                    conversation_data = (
                        await self.memory_manager.get_all_conversation_history(
                            conversation_id
                        )
                    )
                    if conversation_data:
                        for msg in conversation_data:
                            role = msg.get("role", "Unknown").lower()
                            content = msg.get("content", "")
                            if role == "system" or not content.strip():
                                continue
                            formatted_content += f"{content}\n\n"
                            has_content = True
                except Exception as e:
                    self.logger.error(f"Error retrieving from memory manager: {e}")
            return formatted_content if has_content else None
        except Exception as e:
            self.logger.error(f"Error getting conversation history: {e}")
            return None


ExportCommands = EnhancedExportCommands
