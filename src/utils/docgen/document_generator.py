"""
Document Generator for creating professionally formatted PDF and DOCX documents.
"""
import io
import re
from datetime import datetime
from typing import Optional
from typing import cast
from docx.styles.style import ParagraphStyle as DocxParagraphStyle
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from .shared_imports import (
    logging,
    REPORTLAB_AVAILABLE,
    getSampleStyleSheet,
    ParagraphStyle,
    colors,
    A4,
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    ListFlowable,
)
class DocumentGenerator:
    """Utility for generating professionally formatted documents"""
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    async def create_pdf(
        self, content: str, title: Optional[str] = None, author: str = "DeepGem Bot"
    ) -> bytes:
        """Generate a professionally formatted PDF from text content"""
        if not REPORTLAB_AVAILABLE:
            raise ImportError(
                "ReportLab is required for PDF generation. Please install it with: pip install reportlab"
            )
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72,
            title=title or "Generated Document",
            author=author,
        )
        styles = getSampleStyleSheet()
        custom_styles = {
            "CustomTitle": ParagraphStyle(
                name="CustomTitle",
                parent=styles["Heading1"],
                fontSize=24,
                spaceAfter=24,
                textColor=colors.darkblue,
            ),
            "CustomSubtitle": ParagraphStyle(
                name="CustomSubtitle",
                parent=styles["Heading2"],
                fontSize=18,
                spaceBefore=12,
                spaceAfter=6,
                textColor=colors.darkblue,
            ),
            "CustomCode": ParagraphStyle(
                name="CustomCode",
                fontName="Courier",
                fontSize=9,
                backColor=colors.lightgrey,
                leftIndent=36,
                rightIndent=36,
            ),
        }
        for style_name, style in custom_styles.items():
            if style_name not in styles:
                styles.add(style)
        elements = []
        if title:
            elements.append(Paragraph(title, styles["CustomTitle"]))
        else:
            elements.append(Paragraph("Generated Document", styles["CustomTitle"]))
        elements.append(
            Paragraph(
                f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
                styles["Italic"],
            )
        )
        elements.append(Spacer(1, 24))
        for line in content.split("\n"):
            if line.startswith("# "):
                elements.append(Paragraph(line[2:], styles["Heading1"]))
            elif line.startswith("## "):
                elements.append(Paragraph(line[3:], styles["Heading2"]))
            elif line.startswith("### "):
                elements.append(Paragraph(line[4:], styles["Heading3"]))
            elif line.startswith("```"):
                code_content = []
                code_block_index = content.split("\n").index(line)
                for code_line in content.split("\n")[code_block_index + 1 :]:
                    if code_line.strip() == "```":
                        break
                    code_content.append(code_line)
                if code_content:
                    code_text = "\n".join(code_content)
                    elements.append(Paragraph(code_text, styles["CustomCode"]))
                continue
            elif line.startswith("- "):
                bullet_text = Paragraph(line[2:], styles["Normal"])
                elements.append(
                    ListFlowable([bullet_text], bulletType="bullet", start=None)
                )
            elif line.strip().startswith("*") and line.strip()[1:2].isspace():
                bullet_text = Paragraph(line.strip()[2:], styles["Normal"])
                elements.append(
                    ListFlowable([bullet_text], bulletType="bullet", start=None)
                )
            elif line.startswith("> "):
                elements.append(
                    Paragraph(
                        line[2:],
                        ParagraphStyle(
                            name="BlockQuote",
                            parent=styles["Normal"],
                            leftIndent=36,
                            textColor=colors.darkslategray,
                        ),
                    )
                )
            elif line.strip() == "":
                elements.append(Spacer(1, 12))
            else:
                elements.append(Paragraph(line, styles["Normal"]))
        doc.build(elements)
        buffer.seek(0)
        return buffer.getvalue()
    async def create_docx(
        self, content: str, title: Optional[str] = None, author: str = "DeepGem Bot"
    ) -> bytes:
        """Generate a professionally formatted DOCX from text content"""
        content = self._clean_markdown_formatting(content)
        content = re.sub(
            r"^\s*[•\*]\s*\*\*(.*?)\*\*\s*", r"* **\1**", content, flags=re.MULTILINE
        )
        doc = Document()
        buffer = io.BytesIO()
        try:
            style = doc.styles["Normal"]
            style_font = cast(DocxParagraphStyle, style).font
            style_font.name = "Calibri"
            style_font.size = Pt(11)
        except Exception as e:
            self.logger.warning(f"Could not set font for Normal style: {e}")
        for i in range(1, 4):
            try:
                heading_style = doc.styles[f"Heading {i}"]
                heading_font = cast(DocxParagraphStyle, heading_style).font
                heading_font.name = "Calibri"
                heading_font.color.rgb = RGBColor(0, 70, 127)
                if i == 1:
                    heading_font.size = Pt(18)
                    heading_font.bold = True
                elif i == 2:
                    heading_font.size = Pt(16)
                    heading_font.bold = True
                else:
                    heading_font.size = Pt(14)
                    heading_font.bold = True
            except Exception as e:
                self.logger.warning(f"Could not set font for Heading {i} style: {e}")
        doc.core_properties.author = author
        doc.core_properties.title = title or "Generated Document"
        if title:
            doc_title = doc.add_heading(title, 0)
        else:
            doc_title = doc.add_heading("Generated Document", 0)
        doc_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        metadata = doc.add_paragraph(
            f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        )
        run = metadata.runs[0] if metadata.runs else metadata.add_run()
        run.italic = True
        doc.add_paragraph()
        current_section_level = 0
        in_table = False
        current_table = None
        current_paragraph = None
        lines = content.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]
            if line.startswith("# "):
                h1 = doc.add_heading(line[2:], level=1)
                h1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                current_paragraph = None
                current_section_level = 1
                in_table = False
            elif line.startswith("## "):
                h2 = doc.add_heading(line[3:], level=2)
                h2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                current_paragraph = None
                current_section_level = 2
                in_table = False
            elif line.startswith("### "):
                h3 = doc.add_heading(line[4:], level=3)
                h3.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                current_paragraph = None
                current_section_level = 3
                in_table = False
            elif line.strip() == "":
                if not in_table and (
                    not current_paragraph or current_paragraph.text.strip()
                ):
                    doc.add_paragraph()
                    current_paragraph = None
            else:
                pass
                if line.startswith("- ") or line.startswith("* "):
                    p = doc.add_paragraph(style="List Bullet")
                    self._process_mixed_formatting(p, line[2:])
                    current_paragraph = None
                    in_table = False
                elif re.match(r"^\d+\.\s", line):
                    match = re.match(r"^(\d+)\.(\s+)(.+)$", line)
                    if match:
                        num, spaces, content = match.groups()
                        p = doc.add_paragraph(style="List Number")
                        self._process_mixed_formatting(p, content)
                        current_paragraph = None
                    in_table = False
                elif line.startswith("> "):
                    p = doc.add_paragraph(line[2:])
                    p.style = "Intense Quote"
                    current_paragraph = None
                    in_table = False
                elif line.startswith("|") and line.endswith("|"):
                    cells = [cell.strip() for cell in line.strip("|").split("|")]
                    if any(re.match(r"^[-:]+$", cell.strip()) for cell in cells):
                        i += 1
                        continue
                    if not in_table:
                        in_table = True
                        current_table = doc.add_table(rows=0, cols=len(cells))
                        current_table.style = "Table Grid"
                        header_cells = current_table.add_row().cells
                        for j, cell_text in enumerate(cells):
                            cell = header_cells[j]
                            if "**" in cell_text or "*" in cell_text:
                                paragraph = cell.paragraphs[0]
                                paragraph.text = ""
                                self._process_mixed_formatting(paragraph, cell_text)
                            else:
                                cell.text = cell_text
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for cell in header_cells:
                            cell_properties = cell._tc.get_or_add_tcPr()
                            cell_shading = OxmlElement("w:shd")
                            cell_shading.set(qn("w:fill"), "D0D0D0")
                            cell_properties.append(cell_shading)
                        pass
                    else:
                        if current_table is not None:
                            row_cells = current_table.add_row().cells
                            for j, cell_text in enumerate(cells):
                                if j < len(row_cells):
                                    cell = row_cells[j]
                                    if "**" in cell_text or "*" in cell_text:
                                        paragraph = cell.paragraphs[0]
                                        paragraph.text = ""
                                        self._process_mixed_formatting(
                                            paragraph, cell_text
                                        )
                                    else:
                                        cell.text = cell_text
                    i += 1
                    continue
                else:
                    in_table = False
                    if not current_paragraph:
                        current_paragraph = doc.add_paragraph()
                    is_summary_or_conclusion = False
                    if current_section_level == 2:
                        section_title = lines[max(0, i - 2) : i]
                        section_title = [
                            line for line in section_title if line.startswith("## ")
                        ]
                        if section_title and (
                            "summary" in section_title[0].lower()
                            or "conclusion" in section_title[0].lower()
                        ):
                            is_summary_or_conclusion = True
                    if "**" in line or "*" in line:
                        self._process_mixed_formatting(current_paragraph, line)
                    else:
                        if is_summary_or_conclusion:
                            run = current_paragraph.add_run(line)
                            run.italic = True
                        else:
                            current_paragraph.add_run(line)
            i += 1
        if in_table:
            doc.add_paragraph()
        has_summary = any(
            "summary" in line.lower() for line in lines if line.startswith("## ")
        )
        has_conclusion = any(
            "conclusion" in line.lower() for line in lines if line.startswith("## ")
        )
        if not has_summary and not has_conclusion:
            doc.add_heading("Summary", level=2)
            summary_para = doc.add_paragraph()
            summary_para.add_run("This document provides an overview of ").italic = True
            if title:
                title_part = summary_para.add_run(title.lower())
                title_part.italic = True
                title_part.bold = True
            else:
                summary_para.add_run("the topic").italic = True
            summary_para.add_run(
                ". Refer to the sections above for detailed information."
            ).italic = True
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    def _process_mixed_formatting(self, paragraph, text):
        """Process text with mixed formatting (bold, code, italic)"""
        in_bold = False
        in_code = False
        in_italic = False
        buffer = ""
        i = 0
        while i < len(text):
            if text[i : i + 2] == "**":
                if buffer:
                    run = paragraph.add_run(buffer)
                    if in_bold:
                        run.bold = True
                    if in_italic:
                        run.italic = True
                    if in_code:
                        run.font.name = "Courier New"
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(63, 127, 127)
                    buffer = ""
                in_bold = not in_bold
                i += 2
                continue
            elif text[i] == "`":
                if buffer:
                    run = paragraph.add_run(buffer)
                    if in_bold:
                        run.bold = True
                    if in_italic:
                        run.italic = True
                    if in_code:
                        run.font.name = "Courier New"
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(63, 127, 127)
                    buffer = ""
                in_code = not in_code
                i += 1
                continue
            elif text[i] == "*" and not text[i : i + 2] == "**" and not in_code:
                if buffer:
                    run = paragraph.add_run(buffer)
                    if in_bold:
                        run.bold = True
                    if in_italic:
                        run.italic = True
                    buffer = ""
                in_italic = not in_italic
                i += 1
                continue
            else:
                buffer += text[i]
                i += 1
        if buffer:
            run = paragraph.add_run(buffer)
            if in_bold:
                run.bold = True
            if in_italic:
                run.italic = True
            if in_code:
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(63, 127, 127)
    def _process_mixed_bold_italic(self, paragraph, text):
        """Process text with mixed bold and italic formatting"""
        parts = text.split("**")
        for i, part in enumerate(parts):
            if i % 2 == 0:
                if "*" in part:
                    italic_parts = part.split("*")
                    for j, italic_part in enumerate(italic_parts):
                        if j % 2 == 0:
                            paragraph.add_run(italic_part)
                        else:
                            italic_run = paragraph.add_run(italic_part)
                            italic_run.italic = True
                else:
                    paragraph.add_run(part)
            else:
                if "*" in part:
                    italic_parts = part.split("*")
                    for j, italic_part in enumerate(italic_parts):
                        if j % 2 == 0:
                            bold_run = paragraph.add_run(italic_part)
                            bold_run.bold = True
                        else:
                            bold_italic_run = paragraph.add_run(italic_part)
                            bold_italic_run.bold = True
                            bold_italic_run.italic = True
                else:
                    bold_run = paragraph.add_run(part)
                    bold_run.bold = True
    def _clean_markdown_formatting(self, content: str) -> str:
        """Fix common markdown formatting issues"""
        lines = content.splitlines()
        cleaned_lines = []
        for i, line in enumerate(lines):
            if re.match(r"^#{1,6}[^#\s]", line):
                for j in range(6, 0, -1):
                    pattern = r"^#{" + str(j) + r"}([^#\s])"
                    if re.match(pattern, line):
                        line = re.sub(pattern, "#" * j + r" \1", line)
                        break
            if re.match(r"^\s*[\*\-•]([^\s])", line):
                line = re.sub(r"^(\s*[\*\-•])([^\s])", r"\1 \2", line)
            if re.match(r"^\s*\d+\.[^\s]", line):
                line = re.sub(r"^(\s*\d+\.)([^\s])", r"\1 \2", line)
            if line.strip().startswith("•"):
                line = "* " + line.strip()[1:].lstrip()
            line = re.sub(r"\*\*\s+", r"**", line)
            if "**" in line and line.count("**") % 2 != 0:
                if not any(
                    lines[j].strip().startswith("**")
                    for j in range(i + 1, min(i + 3, len(lines)))
                ):
                    line = line + "**"
            if re.match(r"^\s*\|.*\|\s*$", line):
                parts = line.split("|")
                for j in range(len(parts)):
                    parts[j] = parts[j].strip()
                line = "|" + "|".join(f" {p} " for p in parts[1:-1]) + "|"
            cleaned_lines.append(line)
        result_lines = []
        i = 0
        while i < len(cleaned_lines):
            current_line = cleaned_lines[i]
            result_lines.append(current_line)
            if re.match(r"^#{1,6}\s", current_line) and i < len(cleaned_lines) - 1:
                next_non_empty = i + 1
                while (
                    next_non_empty < len(cleaned_lines)
                    and not cleaned_lines[next_non_empty].strip()
                ):
                    next_non_empty += 1
                if next_non_empty < len(cleaned_lines) and re.match(
                    r"^#{1,6}\s", cleaned_lines[next_non_empty]
                ):
                    result_lines.append("")
                    i = next_non_empty - 1
            i += 1
        result = "\n".join(result_lines)
        result = re.sub(r"\n{3,}", "\n\n", result)
        result = re.sub(r"```(\w+)?\s+", r"```\1\n", result)
        result = re.sub(r"\s+```", r"\n```", result)
        result = re.sub(r"^\s*\*\s*([^*])", r"* \1", result, flags=re.MULTILINE)
        return result
