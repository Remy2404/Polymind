import logging
import google.generativeai as genai
import io
import httpx
from typing import Optional, List, Dict, BinaryIO, Union, Any
import asyncio
from utils.telegramlog import telegram_logger
import os
from dotenv import load_dotenv
from telegram import Bot
from services.knowledge_graph import KnowledgeGraph
import mimetypes
import re
import json
import markdown
import tempfile
import sys
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE

# Conditionally import WeasyPrint - it might not work on Windows systems without extra libraries
WEASYPRINT_AVAILABLE = False
try:
    from weasyprint import HTML, CSS

    WEASYPRINT_AVAILABLE = True
except ImportError as e:
    logging.warning(
        f"WeasyPrint import error: {str(e)}. PDF generation will use fallback method."
    )
except Exception as e:
    logging.warning(
        f"WeasyPrint error: {str(e)}. PDF generation will use fallback method."
    )

# Try to import a fallback PDF library (reportlab) for PDF generation
REPORTLAB_AVAILABLE = False
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
        Preformatted,
        Image,
        ListItem,
        ListFlowable,
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch, cm
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY

    REPORTLAB_AVAILABLE = True
except ImportError:
    logging.warning(
        "ReportLab not installed. Run 'pip install reportlab' for better PDF fallback support."
    )

# Load environment variables
load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")


class DocumentGenerator:
    """Utility for generating professionally formatted documents"""

    def __init__(self):
        self.logger = logging.getLogger(__name__)

    async def create_pdf(
        self, content: str, title: str = None, author: str = "DeepGem Bot"
    ) -> bytes:
        """Generate a professionally formatted PDF from text content"""
        buffer = io.BytesIO()

        # Set up the document
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72,
            title=title or "Generated Document",
            author=author,
        )

        # Get basic styles
        styles = getSampleStyleSheet()

        # Check if styles already exist before adding them
        custom_styles = {
            "CustomTitle": ParagraphStyle(
                name="CustomTitle",
                parent=styles["Heading1"],
                fontSize=24,
                spaceAfter=24,
                textColor=colors.darkblue,
            ),
            "CustomSubtitle": ParagraphStyle(
                name="CustomSubtitle",
                parent=styles["Heading2"],
                fontSize=18,
                spaceBefore=12,
                spaceAfter=6,
                textColor=colors.darkblue,
            ),
            "CustomCode": ParagraphStyle(
                name="CustomCode",
                fontName="Courier",
                fontSize=9,
                backColor=colors.lightgrey,
                leftIndent=36,
                rightIndent=36,
            ),
        }

        # Add custom styles that don't conflict with existing names
        for style_name, style in custom_styles.items():
            if style_name not in styles:
                styles.add(style)

        # Process content
        elements = []

        # Add title
        if title:
            elements.append(Paragraph(title, styles["CustomTitle"]))
        else:
            elements.append(Paragraph("Generated Document", styles["CustomTitle"]))

        # Add metadata
        elements.append(
            Paragraph(
                f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
                styles["Italic"],
            )
        )
        elements.append(Spacer(1, 24))

        # Split content and process
        for line in content.split("\n"):
            # Process markdown-style formatting
            if line.startswith("# "):
                # H1 header
                elements.append(Paragraph(line[2:], styles["Heading1"]))
            elif line.startswith("## "):
                # H2 header
                elements.append(Paragraph(line[3:], styles["Heading2"]))
            elif line.startswith("### "):
                # H3 header
                elements.append(Paragraph(line[4:], styles["Heading3"]))
            elif line.startswith("```"):
                # Code block - collect until closing ```
                code_content = []
                in_code_block = True
                code_block_index = content.split("\n").index(line)

                # Skip the opening ``` line
                for code_line in content.split("\n")[code_block_index + 1 :]:
                    if code_line.strip() == "```":
                        in_code_block = False
                        break
                    code_content.append(code_line)

                if code_content:
                    code_text = "\n".join(code_content)
                    elements.append(Paragraph(code_text, styles["CustomCode"]))
                continue
            elif line.startswith("- "):
                # List item with dash - using ListFlowable to wrap the ListItem
                bullet_text = Paragraph(line[2:], styles["Normal"])
                list_item = ListItem(bullet_text, leftIndent=36)
                elements.append(
                    ListFlowable([list_item], bulletType="bullet", start=None)
                )
            elif line.strip().startswith("*") and line.strip()[1:2].isspace():
                # List item with asterisk - using ListFlowable to wrap the ListItem
                bullet_text = Paragraph(line.strip()[2:], styles["Normal"])
                list_item = ListItem(bullet_text, leftIndent=36)
                elements.append(
                    ListFlowable([list_item], bulletType="bullet", start=None)
                )
            elif line.startswith("> "):
                # Blockquote
                elements.append(
                    Paragraph(
                        line[2:],
                        ParagraphStyle(
                            name="BlockQuote",
                            parent=styles["Normal"],
                            leftIndent=36,
                            textColor=colors.darkslategray,
                        ),
                    )
                )
            elif line.strip() == "":
                # Empty line
                elements.append(Spacer(1, 12))
            else:
                # Regular paragraph
                elements.append(Paragraph(line, styles["Normal"]))

        # Build the document
        doc.build(elements)
        buffer.seek(0)
        return buffer.getvalue()

    async def create_docx(
        self, content: str, title: str = None, author: str = "DeepGem Bot"
    ) -> bytes:
        """Generate a professionally formatted DOCX from text content"""
        # First, pre-process content to fix common formatting issues
        content = self._clean_markdown_formatting(content)

        # Special handling for bullet points with bold markers
        content = re.sub(
            r"^\s*[•\*]\s*\*\*(.*?)\*\*\s*", r"* **\1**", content, flags=re.MULTILINE
        )

        doc = Document()
        buffer = io.BytesIO()

        # Apply professional document styling
        # Set default font
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        # Customize heading styles for visual consistency
        for i in range(1, 4):
            heading_style = doc.styles[f"Heading {i}"]
            heading_style.font.name = "Calibri"
            heading_style.font.color.rgb = RGBColor(0, 70, 127)  # Dark blue

            # Different sizes for different heading levels
            if i == 1:
                heading_style.font.size = Pt(18)
                heading_style.font.bold = True
            elif i == 2:
                heading_style.font.size = Pt(16)
                heading_style.font.bold = True
            else:
                heading_style.font.size = Pt(14)
                heading_style.font.bold = True

        # Document properties
        doc.core_properties.author = author
        doc.core_properties.title = title or "Generated Document"

        # Add title
        if title:
            doc_title = doc.add_heading(title, 0)
        else:
            doc_title = doc.add_heading("Generated Document", 0)
        doc_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add metadata
        metadata = doc.add_paragraph(
            f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        )
        metadata.italic = True
        doc.add_paragraph()  # Empty line

        # Track section headings and table state
        current_section_level = 0
        current_section_content = False
        in_table = False
        current_table = None
        table_header_processed = False

        # Process content with markdown-style parsing
        current_paragraph = None
        in_code_block = False

        # Split content by lines for processing
        lines = content.split("\n")
        i = 0

        while i < len(lines):
            line = lines[i]

            # Handle headings and track section structure
            if line.startswith("# "):
                # Main title (H1)
                h1 = doc.add_heading(line[2:], level=1)
                h1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                current_paragraph = None
                current_section_level = 1
                current_section_content = False
                in_table = False  # Exit table mode when new section starts
            elif line.startswith("## "):
                # Section heading (H2)
                h2 = doc.add_heading(line[3:], level=2)
                h2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                current_paragraph = None
                current_section_level = 2
                current_section_content = False
                in_table = False  # Exit table mode when new section starts
            elif line.startswith("### "):
                # Subsection heading (H3)
                h3 = doc.add_heading(line[4:], level=3)
                h3.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                current_paragraph = None
                current_section_level = 3
                current_section_content = False
                in_table = False  # Exit table mode when new section starts
            elif line.strip() == "":
                # Empty line - add paragraph break if needed
                if not in_table and (
                    not current_paragraph or current_paragraph.text.strip()
                ):
                    doc.add_paragraph()
                    current_paragraph = None
            else:
                # Content line - mark that we've added content to the current section
                current_section_content = True

                # Process different content types
                if line.startswith("- ") or line.startswith("* "):
                    # List item with dash or asterisk
                    p = doc.add_paragraph(style="List Bullet")
                    # Use full text processing to handle formatting in list items
                    self._process_mixed_formatting(p, line[2:])
                    current_paragraph = None
                    in_table = False  # Exit table mode for list items
                elif re.match(r"^\d+\.\s", line):
                    # Numbered list item (more flexible matching)
                    match = re.match(r"^(\d+)\.(\s+)(.+)$", line)
                    if match:
                        num, spaces, content = match.groups()
                        p = doc.add_paragraph(style="List Number")
                        # Use full text processing to handle formatting in list items
                        self._process_mixed_formatting(p, content)
                        current_paragraph = None
                    in_table = False  # Exit table mode for list items
                elif line.startswith("> "):
                    # Blockquote
                    p = doc.add_paragraph(line[2:])
                    p.style = "Intense Quote"
                    current_paragraph = None
                    in_table = False  # Exit table mode for blockquotes
                elif line.startswith("|") and line.endswith("|"):
                    # Table row detection
                    cells = [cell.strip() for cell in line.strip("|").split("|")]

                    # Check if this is a separator row
                    if any(re.match(r"^[-:]+$", cell.strip()) for cell in cells):
                        # For separator row, just track alignment info
                        i += 1
                        continue

                    # Determine if this is the first row of a table
                    if not in_table:
                        in_table = True
                        table_header_processed = False
                        current_table = doc.add_table(rows=0, cols=len(cells))
                        current_table.style = "Table Grid"

                        # Apply professional table styling
                        # Make first row as header with special formatting
                        header_cells = current_table.add_row().cells

                        # Add the header content with formatting
                        for j, cell_text in enumerate(cells):
                            cell = header_cells[j]
                            # Process formatting in cell text
                            if "**" in cell_text or "*" in cell_text:
                                paragraph = cell.paragraphs[0]
                                paragraph.text = ""
                                self._process_mixed_formatting(paragraph, cell_text)
                            else:
                                cell.text = cell_text

                            # Apply header formatting
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                        # Apply special styling to header row
                        for cell in header_cells:
                            cell_properties = cell._tc.get_or_add_tcPr()
                            cell_shading = OxmlElement("w:shd")
                            cell_shading.set(
                                qn("w:fill"), "D0D0D0"
                            )  # Light gray background
                            cell_properties.append(cell_shading)

                        table_header_processed = True
                    else:
                        # Add a regular row to the table
                        row_cells = current_table.add_row().cells

                        # Add cell content with formatting
                        for j, cell_text in enumerate(cells):
                            if j < len(row_cells):
                                cell = row_cells[j]
                                # Process formatting in cell text
                                if "**" in cell_text or "*" in cell_text:
                                    paragraph = cell.paragraphs[0]
                                    paragraph.text = ""
                                    self._process_mixed_formatting(paragraph, cell_text)
                                else:
                                    cell.text = cell_text

                    # No need to set current_paragraph since we're in table mode
                    i += 1
                    continue
                else:
                    # Regular paragraph text - exit table mode
                    in_table = False

                    # Create a new paragraph
                    if not current_paragraph:
                        current_paragraph = doc.add_paragraph()

                    # Handle special sections
                    is_summary_or_conclusion = False
                    if current_section_level == 2:
                        section_title = lines[max(0, i - 2) : i]
                        section_title = [
                            l for l in section_title if l.startswith("## ")
                        ]
                        if section_title and (
                            "summary" in section_title[0].lower()
                            or "conclusion" in section_title[0].lower()
                        ):
                            is_summary_or_conclusion = True

                    # Handle formatting in the text (bold, italic)
                    if "**" in line or "*" in line:
                        self._process_mixed_formatting(current_paragraph, line)
                    else:
                        # Apply special formatting for summary/conclusion
                        if is_summary_or_conclusion:
                            run = current_paragraph.add_run(line)
                            run.italic = True
                        else:
                            current_paragraph.add_run(line)

            # Move to next line
            i += 1

        # If we ended in the middle of a table, add extra spacing
        if in_table:
            doc.add_paragraph()

        # Add a summary section if none exists
        has_summary = any(
            "summary" in line.lower() for line in lines if line.startswith("## ")
        )
        has_conclusion = any(
            "conclusion" in line.lower() for line in lines if line.startswith("## ")
        )

        if not has_summary and not has_conclusion:
            # Add a basic summary section
            doc.add_heading("Summary", level=2)
            summary_para = doc.add_paragraph()
            summary_para.add_run("This document provides an overview of ").italic = True
            if title:
                title_part = summary_para.add_run(title.lower())
                title_part.italic = True
                title_part.bold = True
            else:
                summary_para.add_run("the topic").italic = True
            summary_para.add_run(
                ". Refer to the sections above for detailed information."
            ).italic = True

        # Save to buffer
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _process_mixed_formatting(self, paragraph, text):
        """Process text with mixed formatting (bold, code, italic)"""
        # This is a simplified approach - for complex mixed formatting
        # a proper parser would be better, but this handles common cases

        current_pos = 0
        in_bold = False
        in_code = False
        in_italic = False
        buffer = ""

        i = 0
        while i < len(text):
            # Check for bold marker
            if text[i : i + 2] == "**":
                # Flush buffer
                if buffer:
                    run = paragraph.add_run(buffer)
                    if in_bold:
                        run.bold = True
                    if in_italic:
                        run.italic = True
                    if in_code:
                        run.font.name = "Courier New"
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(63, 127, 127)
                    buffer = ""

                in_bold = not in_bold
                i += 2
                continue

            # Check for code marker
            elif text[i] == "`":
                # Flush buffer
                if buffer:
                    run = paragraph.add_run(buffer)
                    if in_bold:
                        run.bold = True
                    if in_italic:
                        run.italic = True
                    if in_code:
                        run.font.name = "Courier New"
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(63, 127, 127)
                    buffer = ""

                in_code = not in_code
                i += 1
                continue

            # Check for italic marker (not inside code)
            elif text[i] == "*" and not text[i : i + 2] == "**" and not in_code:
                # Flush buffer
                if buffer:
                    run = paragraph.add_run(buffer)
                    if in_bold:
                        run.bold = True
                    if in_italic:
                        run.italic = True
                    buffer = ""

                in_italic = not in_italic
                i += 1
                continue

            # Regular character
            else:
                buffer += text[i]
                i += 1

        # Flush any remaining buffer
        if buffer:
            run = paragraph.add_run(buffer)
            if in_bold:
                run.bold = True
            if in_italic:
                run.italic = True
            if in_code:
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(63, 127, 127)

    def _process_mixed_bold_italic(self, paragraph, text):
        """Process text with mixed bold and italic formatting"""
        # First handle bold
        parts = text.split("**")
        for i, part in enumerate(parts):
            if i % 2 == 0:  # Not bold
                # Check for italic in non-bold text
                if "*" in part:
                    italic_parts = part.split("*")
                    for j, italic_part in enumerate(italic_parts):
                        if j % 2 == 0:  # Regular text
                            paragraph.add_run(italic_part)
                        else:  # Italic text
                            italic_run = paragraph.add_run(italic_part)
                            italic_run.italic = True
                else:
                    paragraph.add_run(part)
            else:  # Bold text
                # Check for italic in bold text
                if "*" in part:
                    italic_parts = part.split("*")
                    for j, italic_part in enumerate(italic_parts):
                        if j % 2 == 0:  # Bold only
                            bold_run = paragraph.add_run(italic_part)
                            bold_run.bold = True
                        else:  # Bold and italic
                            bold_italic_run = paragraph.add_run(italic_part)
                            bold_italic_run.bold = True
                            bold_italic_run.italic = True
                else:
                    bold_run = paragraph.add_run(part)
                    bold_run.bold = True

    def _clean_markdown_formatting(self, content: str) -> str:
        """Fix common markdown formatting issues"""
        lines = content.splitlines()
        cleaned_lines = []

        # First pass: Fix basic formatting issues
        for i, line in enumerate(lines):
            # Fix header spacing (ensure space after #)
            if re.match(r"^#{1,6}[^#\s]", line):
                for j in range(6, 0, -1):
                    pattern = r"^#{" + str(j) + r"}([^#\s])"
                    if re.match(pattern, line):
                        line = re.sub(pattern, "#" * j + r" \1", line)
                        break

            # Fix bullet point spacing (ensure space after * or -)
            if re.match(r"^\s*[\*\-•]([^\s])", line):
                line = re.sub(r"^(\s*[\*\-•])([^\s])", r"\1 \2", line)

            # Fix numbered list spacing
            if re.match(r"^\s*\d+\.[^\s]", line):
                line = re.sub(r"^(\s*\d+\.)([^\s])", r"\1 \2", line)

            # Fix improperly formatted bullet points with unicode bullet character
            if line.strip().startswith("•"):
                line = "* " + line.strip()[1:].lstrip()

            # Fix improperly formatted bold/italic text
            # Replace ** followed by a space with just **
            line = re.sub(r"\*\*\s+", r"**", line)
            # Fix missing closing ** for bold text
            if "**" in line and line.count("**") % 2 != 0:
                if not any(
                    lines[j].strip().startswith("**")
                    for j in range(i + 1, min(i + 3, len(lines)))
                ):
                    line = line + "**"

            # Fix table formatting - ensure proper spacing in table cells
            if re.match(r"^\s*\|.*\|\s*$", line):
                # Split by | and process each cell
                parts = line.split("|")
                for j in range(len(parts)):
                    parts[j] = parts[j].strip()
                # Rebuild the line with proper spacing
                line = "|" + "|".join(f" {p} " for p in parts[1:-1]) + "|"

            cleaned_lines.append(line)

        # Second pass: Remove excessive empty lines between headings
        result_lines = []
        i = 0
        while i < len(cleaned_lines):
            current_line = cleaned_lines[i]
            result_lines.append(current_line)

            # Check for consecutive headers with empty lines between them
            if re.match(r"^#{1,6}\s", current_line) and i < len(cleaned_lines) - 1:
                next_non_empty = i + 1
                while (
                    next_non_empty < len(cleaned_lines)
                    and not cleaned_lines[next_non_empty].strip()
                ):
                    next_non_empty += 1

                if next_non_empty < len(cleaned_lines) and re.match(
                    r"^#{1,6}\s", cleaned_lines[next_non_empty]
                ):
                    # Found another header after empty lines - add just one empty line
                    result_lines.append("")
                    i = next_non_empty - 1  # will be incremented in the loop

            i += 1

        # Final join with proper line handling
        result = "\n".join(result_lines)

        # Fix consecutive newlines (max 2)
        result = re.sub(r"\n{3,}", "\n\n", result)

        # Ensure code blocks are properly formatted
        result = re.sub(r"```(\w+)?\s+", r"```\1\n", result)
        result = re.sub(r"\s+```", r"\n```", result)

        # Fix isolated asterisks that might be meant as bullet points
        result = re.sub(r"^\s*\*\s*([^*])", r"* \1", result, flags=re.MULTILINE)

        return result


class DocumentProcessor:
    """Handles document processing using the Gemini API."""

    # Supported MIME types for document processing
    # Update the SUPPORTED_MIME_TYPES dictionary
    SUPPORTED_MIME_TYPES = {
        "pdf": "application/pdf",
        "js": "text/plain",  # Changed to text/plain
        "py": "text/plain",  # Changed to text/plain
        "txt": "text/plain",
        "html": "text/html",
        "css": "text/css",
        "md": "text/markdown",
        "csv": "text/csv",
        "xml": "text/xml",
        "rtf": "text/rtf",
        # Add defaults for common programming languages
        "java": "text/plain",
        "cpp": "text/plain",
        "c": "text/plain",
        "cs": "text/plain",
        "php": "text/plain",
        "rb": "text/plain",
        "go": "text/plain",
        "swift": "text/plain",
        "kt": "text/plain",
        "rs": "text/plain",
        "ts": "text/plain",
        "sql": "text/plain",
        "sh": "text/plain",
        "yaml": "text/plain",
        "json": "text/plain",
    }

    # Code file extensions that can be converted to text
    CODE_EXTENSIONS = {
        "py",
        "js",
        "java",
        "cpp",
        "c",
        "cs",
        "php",
        "rb",
        "go",
        "swift",
        "kt",
        "rs",
        "ts",
        "html",
        "css",
        "sql",
        "sh",
        "yaml",
        "json",
        "xml",
        "md",
    }

    def __init__(self, bot: Bot, knowledge_graph: Optional[KnowledgeGraph] = None):
        """Initialize the DocumentProcessor with Gemini API configuration."""
        self.bot = bot
        self.knowledge_graph = knowledge_graph
        self.logger = logging.getLogger(__name__)

        if not GEMINI_API_KEY:
            raise ValueError("GEMINI_API_KEY not found or empty")

        genai.configure(api_key=GEMINI_API_KEY)
        self.model = genai.GenerativeModel("gemini-2.0-flash")

        self.generation_config = {
            "temperature": 0.7,
            "top_p": 0.8,
            "top_k": 20,
            "max_output_tokens": 4096,
        }

    async def upload_file(self, file: BinaryIO, mime_type: str) -> Dict[str, Any]:
        """Asynchronously upload a file to the Gemini API."""
        try:
            # Read the file data
            file_data = file.read()
            # Return the file content parts that Gemini API expects
            return {"mime_type": mime_type, "data": file_data}
        except Exception as e:
            self.logger.error(f"Error uploading file: {str(e)}")
            raise

    def get_mime_type(self, file_extension: str) -> str:
        """Get the MIME type for a given file extension with fallback to text/plain for code files."""
        mime_type_mapping = {
            "pdf": "application/pdf",
            "doc": "application/msword",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "ppt": "application/vnd.ms-powerpoint",
            "pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
            "xls": "application/vnd.ms-excel",
            "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "txt": "text/plain",
            "csv": "text/csv",
            "html": "text/html",
            "htm": "text/html",
            "json": "application/json",
            "xml": "application/xml",
            "md": "text/markdown",
            "jpg": "image/jpeg",
            "jpeg": "image/jpeg",
            "png": "image/png",
            "gif": "image/gif",
            "svg": "image/svg+xml",
        }

        extension = file_extension.lower()
        if extension in mime_type_mapping:
            return mime_type_mapping[extension]
        else:
            # Try to get the MIME type using the mimetypes module
            mime_type = mimetypes.guess_type(f"file.{extension}")[0]
            if mime_type:
                return mime_type
            else:
                # Default to octet-stream
                return "application/octet-stream"

    async def process_document_from_url(self, document_url: str, prompt: str) -> str:
        """Process a document from a URL using the Gemini API."""
        try:
            # Get file extension from URL
            file_extension = document_url.split(".")[-1]
            mime_type = self.get_mime_type(file_extension)
            if not mime_type:
                raise ValueError(f"Unsupported file type: {file_extension}")

            # Retrieve document from URL
            async with httpx.AsyncClient() as client:
                response = await client.get(document_url)
                response.raise_for_status()
                doc_data = io.BytesIO(response.content)

            # Upload document using File API with mime_type
            uploaded_file = await self.upload_file(file=doc_data, mime_type=mime_type)

            # Generate response
            response = await asyncio.to_thread(
                self.model.generate_content,
                [uploaded_file, prompt],
                generation_config=self.generation_config,
            )

            return response.text
        except Exception as e:
            self.logger.error(f"Error processing document from URL: {str(e)}")
            raise

    def _convert_code_to_text(
        self, file: BinaryIO, file_extension: str
    ) -> tuple[io.BytesIO, str]:
        """Convert code file to text format."""
        try:
            content = file.read().decode("utf-8")

            # Add file extension as a header
            header = f"// File type: {file_extension}\n\n"
            formatted_content = header + content

            # Convert back to BytesIO with text/plain mime type
            text_file = io.BytesIO(formatted_content.encode("utf-8"))
            return text_file, "text/plain"
        except Exception as e:
            self.logger.error(f"Error converting code to text: {str(e)}")
            raise

    async def process_document_from_file(
        self, file: Union[bytes, BinaryIO], file_extension: str, prompt: str
    ) -> str:
        """Process a document from a file object."""
        try:
            # Convert to BytesIO if we received bytes
            if isinstance(file, bytes):
                file_io = io.BytesIO(file)
            else:
                file_io = file

            # Rewind to the beginning
            file_io.seek(0)

            # Get the MIME type
            mime_type = self.get_mime_type(file_extension)

            # Process the file using Gemini API
            uploaded_file = await self.upload_file(file=file_io, mime_type=mime_type)

            # Generate a response using the appropriate model
            model = genai.GenerativeModel("gemini-2.0-flash")
            response = await asyncio.to_thread(
                model.generate_content,
                [uploaded_file, prompt],
                generation_config={
                    "temperature": 0.2,
                    "top_p": 0.95,
                    "top_k": 32,
                    "max_output_tokens": 4096,
                },
            )

            if hasattr(response, "text"):
                return response.text
            else:
                self.logger.error("No text property found in the response")
                return "Failed to process document: No text in response"

        except Exception as e:
            self.logger.error(f"Error processing document: {str(e)}")
            return f"Error processing document: {str(e)}"

    async def process_multiple_documents(
        self, documents: List[Dict], prompt: str
    ) -> str:
        """Process multiple documents simultaneously."""
        try:
            uploaded_docs = []

            for doc in documents:
                if "url" in doc:
                    async with httpx.AsyncClient() as client:
                        response = await client.get(doc["url"])
                        response.raise_for_status()
                        doc_data = io.BytesIO(response.content)
                else:
                    doc_data = doc["file"]

                mime_type = self.get_mime_type(doc["extension"])
                if not mime_type:
                    raise ValueError(f"Unsupported file type: {doc['extension']}")

                uploaded_doc = await self.upload_file(
                    file=doc_data, mime_type=mime_type
                )
                uploaded_docs.append(uploaded_doc)

            contents = uploaded_docs + [prompt]

            response = await asyncio.to_thread(
                self.model.generate_content,
                contents,
                generation_config=self.generation_config,
            )

            if hasattr(response, "text") and response.text:
                return response.text
            elif isinstance(response, str) and response:
                return response
            return "Sorry, I couldn't process the documents."

        except Exception as e:
            self.logger.error(f"Error processing multiple documents: {str(e)}")
            raise

    async def delete_processed_file(self, file_name: str) -> bool:
        """Delete a processed file from Gemini API storage."""
        try:
            await asyncio.to_thread(genai.delete_file, file_name)
            return True
        except Exception as e:
            self.logger.error(f"Error deleting file {file_name}: {str(e)}")
            return False

    async def list_processed_files(self) -> List[str]:
        """List all files currently stored in Gemini API storage."""
        try:
            files = await asyncio.to_thread(genai.list_files)
            return [f.name for f in files]
        except Exception as e:
            self.logger.error(f"Error listing files: {str(e)}")
            return []

    async def process_document_enhanced(
        self,
        file: Union[bytes, BinaryIO],
        file_extension: str,
        prompt: str,
        user_id: Optional[str] = None,
        document_id: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        Process documents with enhanced capabilities, including code analysis.
        """
        try:
            # Create a unique document ID if not provided
            if not document_id:
                document_id = f"doc_{int(asyncio.get_event_loop().time())}"

            # Process the file content properly
            file_io = io.BytesIO(file) if isinstance(file, bytes) else file
            file_io.seek(0)

            # Get mime type
            mime_type = self.get_mime_type(file_extension)

            # Upload to Gemini
            uploaded_file = await self.upload_file(file=file_io, mime_type=mime_type)

            # Enhanced processing: Generate summary, extract entities, and answer the prompt
            # Create a more comprehensive prompt that captures multiple aspects
            enhanced_prompt = f"""
            Analyze this document with the following approach:
            1. Understand the document layout, structure, and content type
            2. Extract key information from any tables, charts, or structured data
            3. Identify main topics, entities, and important concepts
            
            Then provide:
            - A comprehensive summary of the document's main points
            - Key entities (people, organizations, locations, technologies mentioned)
            - Important dates or time-related information
            - Main topics and concepts covered
            
            Finally, address the user's specific request: {prompt}
            
            Format your response in clear sections with markdown formatting.
            """

            # Use Gemini 1.5 Pro for best document processing results
            pro_model = genai.GenerativeModel("gemini-2.0-flash")

            # Generate response with enhanced configuration
            response = await asyncio.to_thread(
                pro_model.generate_content,
                [uploaded_file, {"text": enhanced_prompt}],
                generation_config={
                    "temperature": 0.2,  # Lower for more factual responses
                    "top_p": 0.95,
                    "top_k": 64,
                    "max_output_tokens": 8192,  # Allow longer responses for document analysis
                    "response_mime_type": "text/plain",
                },
            )

            # Extract text from response
            if not hasattr(response, "text") or not response.text:
                self.logger.error("No text found in the Gemini API response")
                return {
                    "summary": "Error: Failed to analyze document",
                    "entities": {},
                    "text": "Could not process the document",
                }

            response_text = response.text

            # Extract entities for knowledge graph
            entities_result = await self.extract_document_entities(
                file_io, file_extension
            )

            # Add to knowledge graph if available
            knowledge_graph_summary = None
            if self.knowledge_graph and user_id:
                try:
                    file_io.seek(0)  # Rewind file pointer
                    # Extract text for knowledge graph (simpler approach than full OCR)
                    text_model = genai.GenerativeModel("gemini-1.0-pro")
                    text_extraction = await asyncio.to_thread(
                        text_model.generate_content,
                        [uploaded_file, "Extract all text content from this document."],
                        generation_config={
                            "temperature": 0.1,
                            "max_output_tokens": 16384,
                        },
                    )

                    extracted_text = (
                        text_extraction.text if hasattr(text_extraction, "text") else ""
                    )

                    # Add to knowledge graph
                    knowledge_graph_summary = (
                        await self.knowledge_graph.add_document_entities(
                            document_id=document_id,
                            document_content=extracted_text,
                            user_id=user_id,
                        )
                    )

                    self.logger.info(
                        f"Added document {document_id} to knowledge graph for user {user_id}"
                    )
                except Exception as kg_error:
                    self.logger.error(
                        f"Error adding to knowledge graph: {str(kg_error)}"
                    )

            # Prepare response object with all the extracted information
            result = {
                "text": response_text,
                "entities": (
                    entities_result.get("entities", {}) if entities_result else {}
                ),
                "document_id": document_id,
                "knowledge_graph": knowledge_graph_summary,
            }

            return result

        except Exception as e:
            self.logger.error(f"Error in enhanced document processing: {str(e)}")
            return {
                "text": f"Error processing document: {str(e)}",
                "entities": {},
                "document_id": document_id if document_id else "unknown",
            }

    async def extract_document_entities(
        self, file_content: Union[bytes, BinaryIO], file_extension: str
    ) -> Dict[str, Any]:
        """Extract named entities, key topics, and structured data from documents"""
        try:
            # Ensure we have a BytesIO object
            if isinstance(file_content, bytes):
                file_io = io.BytesIO(file_content)
            else:
                file_io = file_content

            # Reset file pointer
            file_io.seek(0)

            # Get mime type
            mime_type = self.get_mime_type(file_extension)

            # Create specialized prompt for entity extraction
            entity_extraction_prompt = """
            Extract the following from this document:
            1. Named entities (people, organizations, locations)
            2. Key topics and themes
            3. Dates and time references
            4. Structured data (tables, lists)
            5. Technologies or products mentioned
            6. Industry-specific terminology
            
            Format your response as JSON with these categories.
            Example format:
            {
              "entities": {
                "people": ["Name1", "Name2"],
                "organizations": ["Org1", "Org2"],
                "locations": ["Location1", "Location2"],
                "dates": ["Date1", "Date2"],
                "technologies": ["Tech1", "Tech2"],
                "concepts": ["Concept1", "Concept2"]
              },
              "topics": ["Topic1", "Topic2"],
              "structured_data": {
                "tables": [{
                  "description": "Table description",
                  "rows": 5,
                  "columns": 3
                }],
                "lists": [{
                  "description": "List description",
                  "items": 6
                }]
              }
            }
            """

            # Upload file to Gemini
            uploaded_file = await self.upload_file(file=file_io, mime_type=mime_type)

            # Use Gemini for extraction with JSON output
            pro_model = genai.GenerativeModel("gemini-2.0-flash")

            # Configure response to be in JSON format
            response = await asyncio.to_thread(
                pro_model.generate_content,
                [uploaded_file, entity_extraction_prompt],
                generation_config={
                    "temperature": 0.1,  # Lower for more deterministic output
                    "response_mime_type": "application/json",
                },
            )

            # Process response
            if hasattr(response, "text"):
                try:
                    # Extract JSON data from response
                    match = re.search(
                        r"```json\s*(.*?)\s*```", response.text, re.DOTALL
                    )
                    if match:
                        json_text = match.group(1)
                        return json.loads(json_text)
                    else:
                        # Try parsing the whole text as JSON
                        return json.loads(response.text)
                except json.JSONDecodeError:
                    # Fallback to minimal entity structure
                    return {
                        "entities": {
                            "people": [],
                            "organizations": [],
                            "locations": [],
                            "dates": [],
                            "technologies": [],
                            "concepts": [],
                        },
                        "raw_extraction": response.text,
                        "format_error": "Could not parse JSON",
                    }

            return {"error": "No extraction results"}

        except Exception as e:
            self.logger.error(f"Entity extraction error: {str(e)}")
            return {"error": str(e)}

    # Add this method to provide language-specific guidance
    def _get_language_specific_prompt(
        self, file_extension: str, base_prompt: str
    ) -> str:
        """Get language-specific analysis prompts based on file extension."""
        language_prompts = {
            "py": """
                Python-specific analysis:
                - Check for PEP 8 compliance
                - Identify use of common libraries (NumPy, Pandas, Flask, etc.)
                - Look for Pythonic patterns and anti-patterns
            """,
            "js": """
                JavaScript-specific analysis:
                - Identify the framework if any (React, Vue, Angular, etc.)
                - Check for modern JS features and patterns
                - Look for potential performance issues
            """,
            "java": """
                Java-specific analysis:
                - Identify design patterns used
                - Check for proper exception handling
                - Review class structure and inheritance
            """,
            "cpp": """
                C++ specific analysis:
                - Check for memory management issues
                - Look for efficient use of STL
                - Identify potential performance bottlenecks
            """,
        }

        language_specific = language_prompts.get(file_extension.lower(), "")
        if language_specific:
            return f"{base_prompt}\n\n{language_specific}"
        return base_prompt
